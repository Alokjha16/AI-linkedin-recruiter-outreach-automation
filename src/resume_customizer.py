import json
import re
from pathlib import Path

from docx import Document
from docx2pdf import convert


BASE_DOCX = Path("templates/Santosh_Anoushika_Komala_Resume.docx")


BASE_SKILLS = [
    "Cybersecurity", "SOC Operations", "SIEM Log Analysis", "Security Monitoring",
    "Incident Response", "Threat Hunting", "Threat Intelligence", "IoC Analysis",
    "Incident Documentation", "SOC Runbooks", "OSINT", "Vulnerability Identification",
    "Identity & Access Management", "Active Directory", "RBAC", "Access Provisioning",
    "Access Validation", "NIST", "MITRE ATT&CK", "Python", "Wireshark", "Burp Suite",
    "Network Fundamentals", "Technical Documentation"
]


JD_KEYWORD_SKILLS = {
    "wazuh": "Wazuh SIEM/XDR Exposure",
    "siem": "SIEM Log Analysis",
    "xdr": "XDR Security Monitoring",
    "edr": "EDR Security Monitoring",
    "crowdstrike": "CrowdStrike Exposure",
    "microsoft defender": "Microsoft Defender Exposure",
    "defender": "Microsoft Defender Exposure",
    "sentinel": "Microsoft Sentinel Exposure",
    "splunk": "Splunk Exposure",
    "qradar": "IBM QRadar Exposure",
    "soc": "SOC Operations",
    "soc l1": "SOC L1/L2 Operations",
    "soc l2": "SOC L2 Operations",
    "security operations center": "Security Operations Center",
    "incident response": "Incident Response",
    "incident response procedures": "Incident Response Procedures",
    "security incidents": "Security Incident Investigation",
    "escalated alerts": "Escalated Alert Investigation",
    "alert": "Alert Triage",
    "threat hunting": "Threat Hunting",
    "threat intelligence": "Threat Intelligence",
    "compromise assessment": "Compromise Assessment",
    "phishing": "Phishing Analysis",
    "malware": "Malware Analysis",
    "advanced cyber threats": "Advanced Threat Analysis",
    "ioc": "IoC Analysis",
    "mitre": "MITRE ATT&CK",
    "nist": "NIST",
    "runbook": "SOC Runbooks",
    "playbook": "SOC Playbooks",
    "active directory": "Active Directory",
    "windows": "Windows Security",
    "linux": "Linux Security",
    "firewall": "Firewall Concepts",
    "cloud security": "Cloud Security Concepts",
    "iam": "Identity & Access Management",
    "rbac": "RBAC",
    "vulnerability": "Vulnerability Identification",
    "osint": "OSINT",
    "python": "Python",
    "wireshark": "Wireshark",
    "burp": "Burp Suite",
    "ceh": "CEH Exposure",
    "cissp": "CISSP Concepts",
    "cism": "CISM Concepts",
    "gcih": "GCIH Concepts",
    "sc-200": "Microsoft SC-200 Exposure",
}


ROLE_RULES = [
    {
        "title": "SOC SECURITY ANALYST",
        "triggers": ["soc", "soc analyst", "soc l1", "soc l2", "security operations center", "siem", "alert triage"],
        "priority": [
            "SOC Operations", "SOC L1/L2 Operations", "SIEM Log Analysis", "Security Monitoring",
            "Alert Triage", "Escalated Alert Investigation", "Incident Response",
            "Security Incident Investigation", "Threat Hunting", "IoC Analysis",
            "SOC Runbooks", "SOC Playbooks", "MITRE ATT&CK", "NIST"
        ],
    },
    {
        "title": "INCIDENT RESPONSE ANALYST",
        "triggers": ["incident response", "incident responder", "security incident", "incident handling"],
        "priority": [
            "Incident Response", "Security Incident Investigation", "Incident Documentation",
            "Alert Triage", "IoC Analysis", "Threat Hunting", "SIEM Log Analysis",
            "SOC Runbooks", "MITRE ATT&CK", "Technical Documentation"
        ],
    },
    {
        "title": "THREAT INTELLIGENCE ANALYST",
        "triggers": ["threat intelligence", "threat hunting", "osint", "ioc", "compromise assessment"],
        "priority": [
            "Threat Intelligence", "Threat Hunting", "OSINT", "IoC Analysis",
            "Compromise Assessment", "Advanced Threat Analysis", "MITRE ATT&CK",
            "Security Reporting", "Incident Response"
        ],
    },
    {
        "title": "IAM SECURITY ANALYST",
        "triggers": ["iam", "identity access", "identity & access", "active directory", "rbac", "access provisioning"],
        "priority": [
            "Identity & Access Management", "Active Directory", "RBAC",
            "Access Provisioning", "Access Validation", "User Access Reviews",
            "Permission Auditing", "IAM Controls", "Technical Documentation"
        ],
    },
    {
        "title": "VULNERABILITY MANAGEMENT ANALYST",
        "triggers": ["vulnerability", "vulnerability management", "security assessment", "risk assessment"],
        "priority": [
            "Vulnerability Identification", "Vulnerability Management", "Security Assessments",
            "Risk Assessment", "Burp Suite", "Wireshark", "Network Fundamentals",
            "Technical Documentation", "NIST"
        ],
    },
    {
        "title": "CLOUD SECURITY ANALYST",
        "triggers": ["cloud security", "aws", "azure", "gcp", "cloud"],
        "priority": [
            "Cloud Security Concepts", "AWS Security Exposure", "Azure Security Exposure",
            "GCP Security Exposure", "Security Assessments", "Authentication Controls",
            "Data Protection", "Incident Response"
        ],
    },
]


def safe_filename(text):
    text = str(text or "").strip()
    text = re.sub(r"[^A-Za-z0-9]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text or "Candidate"


def clean_text(text):
    return re.sub(r"\s+", " ", str(text or "")).strip()


def load_candidate_data():
    path = Path("candidate_data.json")
    if not path.exists():
        return {
            "full_name": "Santosh Anoushika Komala",
            "name": "Santosh Anoushika Komala",
            "title": "Cybersecurity Analyst",
            "email": "santoshanoushikakomala@gmail.com",
            "phone": "+1 203 809 9717",
            "linkedin": "https://www.linkedin.com/in/anoushika-komala/",
            "availability": "M-F: 11 AM - 6 PM EST",
            "location": "West Haven, CT 06516",
            "current_location": "West Haven, CT 06516",
            "work_authorization": "F1 OPT",
            "total_experience": "2 years",
            "open_to_relocate": "Yes",
        }

    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_candidate_value(candidate, *keys, default=""):
    for key in keys:
        value = candidate.get(key)
        if value:
            value = str(value).strip()
            if value.lower() not in ["none", "nan", "n/a", "null", "not mentioned"]:
                return value
    return default


def detect_title(post_text):
    text = clean_text(post_text).lower()
    best_title = "CYBERSECURITY ANALYST"
    best_score = 0

    for rule in ROLE_RULES:
        score = sum(1 for trigger in rule["triggers"] if trigger in text)
        if score > best_score:
            best_score = score
            best_title = rule["title"]

    return best_title


def get_role_priority_skills(title):
    for rule in ROLE_RULES:
        if rule["title"] == title:
            return rule["priority"]
    return []


def extract_jd_skills(post_text, title):
    text = clean_text(post_text).lower()
    skills = []

    skills.extend(get_role_priority_skills(title))

    for keyword, skill in JD_KEYWORD_SKILLS.items():
        if keyword in text:
            skills.append(skill)

    skills.extend(BASE_SKILLS)

    final = []
    seen = set()

    for skill in skills:
        key = skill.lower()
        if key not in seen:
            seen.add(key)
            final.append(skill)

    return final[:26]


def jd_has(post_text, terms):
    text = clean_text(post_text).lower()
    return any(term.lower() in text for term in terms)


def build_summary(title, post_text):
    text = clean_text(post_text).lower()

    if title == "SOC SECURITY ANALYST":
        opening = (
            "Cybersecurity Analyst with hands-on experience in SOC operations, SIEM monitoring, security alert triage, "
            "incident response, threat hunting, security event analysis, and incident documentation."
        )
    elif title == "INCIDENT RESPONSE ANALYST":
        opening = (
            "Cybersecurity Analyst with hands-on experience in incident response, security monitoring, SIEM investigation, "
            "incident documentation, IoC analysis, access anomaly review, and escalation support."
        )
    elif title == "THREAT INTELLIGENCE ANALYST":
        opening = (
            "Cybersecurity Analyst with hands-on experience in threat intelligence, OSINT investigations, threat hunting, "
            "IoC analysis, fraud infrastructure research, scam mapping, and adversary behavior analysis."
        )
    elif title == "IAM SECURITY ANALYST":
        opening = (
            "Cybersecurity Analyst with practical experience in Identity & Access Management, Active Directory, RBAC, "
            "account provisioning, access validation, password management, and permission auditing."
        )
    elif title == "VULNERABILITY MANAGEMENT ANALYST":
        opening = (
            "Cybersecurity Analyst with experience in vulnerability identification, security assessments, risk awareness, "
            "web security testing, technical documentation, and remediation-focused reporting."
        )
    elif title == "CLOUD SECURITY ANALYST":
        opening = (
            "Cybersecurity Analyst with experience in cloud security concepts, authentication controls, data protection, "
            "security assessments, AI-assisted security analytics, and cyber defense practices."
        )
    else:
        opening = (
            "Cybersecurity Analyst with hands-on experience in SOC operations, SIEM log analysis, incident documentation, "
            "threat intelligence, OSINT, vulnerability identification, IAM, Active Directory, security monitoring, and security assessments."
        )

    jd_lines = []

    if jd_has(text, ["wazuh"]):
        jd_lines.append("familiarity with Wazuh SIEM/XDR concepts")
    if jd_has(text, ["microsoft defender", "defender"]):
        jd_lines.append("exposure to Microsoft Defender and endpoint security workflows")
    if jd_has(text, ["crowdstrike"]):
        jd_lines.append("exposure to CrowdStrike and EDR/XDR monitoring concepts")
    if jd_has(text, ["windows", "linux", "active directory", "firewalls", "firewall"]):
        jd_lines.append("working knowledge of Windows, Linux, Active Directory, and firewall concepts")
    if jd_has(text, ["playbook", "runbook"]):
        jd_lines.append("experience documenting SOPs, SOC runbooks, and incident response procedures")
    if jd_has(text, ["mitre", "nist"]):
        jd_lines.append("understanding of MITRE ATT&CK and NIST-aligned security practices")
    if jd_has(text, ["phishing", "malware", "advanced cyber threats"]):
        jd_lines.append("ability to analyze phishing, malware, and advanced cyber threat indicators")
    if jd_has(text, ["cloud security", "cloud"]):
        jd_lines.append("cloud security awareness across modern enterprise environments")

    if jd_lines:
        alignment = " Also brings " + ", ".join(jd_lines[:5]) + "."
    else:
        alignment = ""

    closing = (
        " Skilled in detecting and documenting threats, analyzing security events, supporting incident response workflows, "
        "and preparing clear technical reports for cyber defense and resilience programs."
    )

    return opening + alignment + closing


def build_skills(title, post_text):
    skills = extract_jd_skills(post_text, title)
    return " • ".join(skills)


def replace_paragraph_text_keep_style(paragraph, new_text):
    if paragraph.runs:
        first_run = paragraph.runs[0]

        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        italic = first_run.italic
        underline = first_run.underline

        first_run.text = str(new_text)

        first_run.font.name = font_name or "Arial"
        first_run.font.size = font_size
        first_run.bold = bold
        first_run.italic = italic
        first_run.underline = underline

        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(str(new_text))
        run.font.name = "Arial"


def replace_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        original_text = paragraph.text
        if not original_text:
            continue

        updated_text = original_text
        changed = False

        for key, value in replacements.items():
            if key in updated_text:
                updated_text = updated_text.replace(key, str(value))
                changed = True

        if changed:
            replace_paragraph_text_keep_style(paragraph, updated_text)


def replace_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs, replacements)
                replace_in_tables(cell.tables, replacements)


def replace_text_in_docx(doc, replacements):
    replace_in_paragraphs(doc.paragraphs, replacements)
    replace_in_tables(doc.tables, replacements)

    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs, replacements)
        replace_in_tables(section.header.tables, replacements)
        replace_in_paragraphs(section.footer.paragraphs, replacements)
        replace_in_tables(section.footer.tables, replacements)


def check_unreplaced_placeholders(doc):
    texts = []

    for paragraph in doc.paragraphs:
        texts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    texts.append(paragraph.text)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            texts.append(paragraph.text)

        for paragraph in section.footer.paragraphs:
            texts.append(paragraph.text)

    joined = "\n".join(texts)
    return re.findall(r"\{\{[^}]+\}\}", joined)


def customize_resume(post_text, lead_index):
    output_dir = Path("outputs/resumes")
    output_dir.mkdir(parents=True, exist_ok=True)

    if not BASE_DOCX.exists():
        raise FileNotFoundError(f"Resume template not found: {BASE_DOCX}")

    candidate = load_candidate_data()

    name = get_candidate_value(candidate, "name", "full_name", default="Santosh Anoushika Komala")
    safe_name = safe_filename(name)

    custom_docx = output_dir / f"{safe_name}_Resume_{lead_index:03d}.docx"
    custom_pdf = output_dir / f"{safe_name}_Resume_{lead_index:03d}.pdf"

    email = get_candidate_value(candidate, "email", default="santoshanoushikakomala@gmail.com")
    phone = get_candidate_value(candidate, "phone", "mobile", default="+1 203 809 9717")
    location = get_candidate_value(candidate, "location", "current_location", default="West Haven, CT 06516")
    linkedin = get_candidate_value(
        candidate,
        "linkedin",
        "linkedin_url",
        default="https://www.linkedin.com/in/anoushika-komala/",
    )

    title = detect_title(post_text)
    summary = build_summary(title, post_text)
    skills = build_skills(title, post_text)

    doc = Document(BASE_DOCX)

    replacements = {
        "{{NAME}}": name,
        "{{TITLE}}": title,
        "{{EMAIL}}": email,
        "{{PHONE}}": phone,
        "{{LOCATION}}": location,
        "{{LINKEDIN}}": linkedin,
        "{{SUMMARY}}": summary,
        "{{SKILLS}}": skills,
        "{{CORE_COMPETENCIES}}": skills,
        "{{PROJECT_SUMMARY_1}}": "",
        "{{PROJECT_SUMMARY_2}}": "",
        "{{PROJECT_SUMMARY_3}}": "",
    }

    replace_text_in_docx(doc, replacements)

    leftover = check_unreplaced_placeholders(doc)
    if leftover:
        print(f"⚠️ Warning: Unreplaced placeholders found: {leftover}")

    doc.save(custom_docx)

    try:
        convert(str(custom_docx), str(custom_pdf))
    except Exception as e:
        raise RuntimeError(
            f"DOCX created but PDF conversion failed. DOCX path: {custom_docx}. Error: {e}"
        )

    if not custom_pdf.exists():
        raise FileNotFoundError(f"PDF was not created: {custom_pdf}")

    return str(custom_pdf)